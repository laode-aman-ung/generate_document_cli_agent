"""
generator.py — Generate konten tiap section via AI, lalu build .docx
"""

import asyncio
from openai import AsyncOpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

from config import DEEPSEEK_API_KEY, DEEPSEEK_MODEL, OUTPUT_DIR
from services.prompts import get_prompt
from services.document_types import DOCUMENT_TYPES


async def _call_ai(client: AsyncOpenAI, prompt: str) -> str:
    response = await client.chat.completions.create(
        model=DEEPSEEK_MODEL,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=2048,
        temperature=0.7,
    )
    return response.choices[0].message.content.strip()


def _build_docx(doc_type: str, topic: str, sections: dict[str, str]) -> str:
    """Bangun file .docx dari konten section yang sudah di-generate."""
    doc = Document()

    # ── Cover / Judul ──────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(doc_type.upper())
    run.bold = True
    run.font.size = Pt(16)

    topic_para = doc.add_paragraph()
    topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topic_run = topic_para.add_run(topic)
    topic_run.font.size = Pt(13)
    topic_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()  # spasi

    # ── Setiap section ─────────────────────────────────────────────────────
    for section_name, content in sections.items():
        # Heading section
        heading = doc.add_heading(section_name.upper(), level=1)
        heading.runs[0].font.size = Pt(13)

        # Konten
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("## ") or line.startswith("# "):
                sub = doc.add_heading(line.lstrip("#").strip(), level=2)
                sub.runs[0].font.size = Pt(11)
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                r = p.add_run(line.strip("**"))
                r.bold = True
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line[0].isdigit() and ". " in line[:4]:
                doc.add_paragraph(line, style="List Number")
            else:
                doc.add_paragraph(line)

        doc.add_paragraph()  # spasi antar section

    # ── Simpan file ────────────────────────────────────────────────────────
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    safe_topic = "".join(c if c.isalnum() or c in " _-" else "" for c in topic)[:40]
    filename = f"{doc_type.replace(' ', '_')}_{safe_topic.replace(' ', '_')}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath


async def generate_document(
    doc_type: str,
    topic: str,
    language: str,
    citation_style: str,
    progress_callback=None,
) -> str:
    """
    Generate dokumen lengkap section per section.
    progress_callback(section_name, idx, total) dipanggil setiap section selesai.
    Returns path file .docx yang dihasilkan.
    """
    if not DEEPSEEK_API_KEY:
        raise RuntimeError("DEEPSEEK_API_KEY belum diset di .env")

    client = AsyncOpenAI(
        api_key=DEEPSEEK_API_KEY,
        base_url="https://api.deepseek.com",
    )

    sections_list = DOCUMENT_TYPES[doc_type]
    generated: dict[str, str] = {}

    for idx, section in enumerate(sections_list, start=1):
        if progress_callback:
            progress_callback(section, idx, len(sections_list))

        prompt = get_prompt(doc_type, section, topic, language, citation_style)
        content = await _call_ai(client, prompt)
        generated[section] = content

    await client.close()

    filepath = _build_docx(doc_type, topic, generated)
    return filepath
